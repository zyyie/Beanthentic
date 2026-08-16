"""
AI Engine for IPOPHL GI Document Analysis.

This module provides text extraction and analysis capabilities for
IPOPHL Geographical Indication registration documents using both
rule-based and machine learning approaches.
"""

import hashlib
import logging
import re
import sys
import uuid
import json
from pathlib import Path
from typing import Dict, List, Tuple

# Support imports whether loaded as `machinelearning.ai_engine` or `ai_engine`
_PROJECT_ROOT = Path(__file__).resolve().parents[1]
_ML_PACKAGE_DIR = Path(__file__).resolve().parent
for _path in (_PROJECT_ROOT, _ML_PACKAGE_DIR):
    if str(_path) not in sys.path:
        sys.path.insert(0, str(_path))

# Text extraction libraries
try:
    import fitz  # PyMuPDF
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    logging.warning("PyMuPDF not available, PDF processing will be limited")

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("python-docx not available, Word document processing will be limited")

try:
    import pytesseract
    from PIL import Image
    import io
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False
    logging.warning("Tesseract OCR not available, scanned PDF processing will be limited")

# ML libraries
try:
    import joblib
    import pandas as pd
    import numpy as np
    import shap
    import matplotlib.pyplot as plt
    ML_AVAILABLE = True
except ImportError:
    ML_AVAILABLE = False
    logging.warning("ML libraries (shap, pandas, numpy, etc.) not available, using rule-based analysis")

class GIAnalyzer:
    """AI Engine for IPOPHL GI Registration and Farmer Readiness Analysis"""

    def __init__(self, uploads_dir: str | None = None, *, auto_train: bool = True):
        self.ml_dir = Path(__file__).resolve().parent
        self.uploads_dir = Path(uploads_dir) if uploads_dir else (self.ml_dir / "uploads")
        self.uploads_dir.mkdir(parents=True, exist_ok=True)

        # Comprehensive checklist of GI-related terms for extraction
        self.gi_checklist = {
            "mandatory_terms": [
                "MoP", "Manual of Specifications", "Causal Link", "Production Process", 
                "Quality Control", "Labeling Rules", "Applicant Entity", "Producers Organization", 
                "Official Receipt", "Application Fee", "Registrability", "Publication",
                "Opposition", "Technical Validation", "Geographical Indication", "GI",
                "Lipa City", "Batangas", "Barako", "Coffee"
            ],
            "optional_terms": [
                "Geographical Area", "LGU", "Organization Documents", "Certificates", "Bylaws",
                "Membership List", "Meeting Minutes", "Attendance Sheets", "Agreement Documents",
                "Territorial Boundaries", "Soil Composition", "Climate Factors", "Historical Reputation",
                "Traditional Knowledge", "Processing Method", "Packaging", "Governing Board",
                "Third-party Observation", "Foreign Protection", "Prior Use", "Distinctive Quality",
                "Flavor Profile", "Aroma", "Roasting Process", "Farming Practices"
            ],
            "term_weights": {
                "Manual of Specifications": 1.5,
                "Causal Link": 2.0,
                "Geographical Indication": 1.5,
                "Lipa City": 1.2,
                "Batangas": 1.2,
                "Barako": 1.2,
                "Quality Control": 1.3,
                "Production Process": 1.3
            }
        }

        # Task-specific checklists aligned to MoP drafting package (7 document groups)
        self.task_checklists = {
            # Phase 1: Justification for the Request for Protection
            "phase1-introduction": {
                "mandatory": ["Kapeng Barako", "Liberica", "Batangas", "Reputation", "Geographical Indication"],
                "optional": ["Lipa", "Tradition", "Distinctive", "Coffea liberica", "Origin"],
            },
            "phase1-history": {
                "mandatory": ["History", "Kapeng Barako", "Batangas", "Heritage"],
                "optional": ["Colonial", "Spanish", "BaCoFFed", "Farmers Federation", "Tradition"],
            },
            "phase1-physical-link": {
                "mandatory": ["Soil", "Climate", "Elevation", "Causal Link", "Territory"],
                "optional": ["Volcanic", "Rainfall", "Taal", "Microclimate", "Physiography"],
            },
            # Phase 2: Technical Part
            "phase2-general": {
                "mandatory": ["Morphological", "Liberica", "Cherry", "Bean", "Sensory"],
                "optional": ["Harvesting Index", "Genetic", "Leaf", "Almond", "Plant Height"],
            },
            "phase2-specific": {
                "mandatory": ["Roasted Coffee", "Ground Coffee", "Degree of Roast", "Aroma"],
                "optional": ["Light Roast", "Medium Roast", "Dark Roast", "Acidity", "Body"],
            },
            "phase2-production": {
                "mandatory": ["Production Process", "Harvesting", "Post-Harvest", "Roasting"],
                "optional": ["Seedling", "Transplanting", "Pruning", "Pulping", "Drying", "NSIC"],
            },
            # Phase 3: Control & Traceability (labelling excluded)
            "phase3-control": {
                "mandatory": ["Internal Control", "Traceability", "BaCoFFed", "Records"],
                "optional": ["Technical Working Group", "Code of Practice", "Certificate of Locality", "LGU"],
            },
            # Legacy procedural zones (older uploads)
            "phase1-product": {
                "mandatory": ["Lipa Barako coffee", "Flavor Profile", "Geographical Origin", "Distinctive Quality"],
                "optional": ["Product Photos", "Aroma", "Roasting Process", "Farming Practices"]
            },
            "phase1-entity": {
                "mandatory": ["Applicant Entity", "Producers Organization", "Legal Standing", "Membership List"],
                "optional": ["LGU", "Certificates", "Bylaws", "Organization Structure"]
            },
            "phase1-stakeholders": {
                "mandatory": ["Stakeholder Consultations", "Meeting Minutes", "Consensus", "Governance Board"],
                "optional": ["Attendance Sheets", "Agreement Documents", "Industry Groups", "Community Support"]
            },
            "phase2-mop": {
                "mandatory": ["Manual of Specifications", "Causal Link", "Production Process", "Quality Control"],
                "optional": ["Geographical Area", "Territorial Boundaries", "Soil Composition", "Climate Factors"]
            },
            "phase2-cert": {
                "mandatory": ["Technical Validation", "Government Certification", "Independent Verification"],
                "optional": ["Foreign Protection", "Proof of Foreign Registration", "Prior Use Evidence"]
            },
            "phase2-details": {
                "mandatory": ["Application Form", "Applicant Name", "Domicile", "Industrial Establishment"],
                "optional": ["Representative Designation", "Commercial Establishment", "Contact Details"]
            },
            "phase3-filing": {
                "mandatory": ["File Application", "Bureau of Trademarks", "Application Package", "Cover Letter"],
                "optional": ["Submission Receipt", "Acknowledgment", "Tracking Number"]
            },
            "phase3-payment": {
                "mandatory": ["Official Receipt", "Application Fee", "Proof of Payment"],
                "optional": ["Exemption Certificate", "Bank Transfer Confirmation", "Payment Date"]
            },
            "phase4-exam": {
                "mandatory": ["Formality Examination", "Substantive Examination", "IP Code Compliance"],
                "optional": ["Examination Reports", "Clarifications", "Technical Responses"]
            },
            "phase4-response": {
                "mandatory": ["Deficiency Notice Response", "Timeframe Compliance", "Corrective Actions"],
                "optional": ["Extensions", "Additional Evidence", "Revised Documents"]
            },
            "phase4-pub": {
                "mandatory": ["Publication for Opposition", "Public Notice Period", "Opposition Period"],
                "optional": ["Third-party Observations", "Opposition Filings", "Response to Objections"]
            },
            "phase5-cert": {
                "mandatory": ["GI Registration Certificate", "Official Notice of Registration", "Registration Number"],
                "optional": ["Award Ceremony", "Public Announcement", "Marketing Materials"]
            },
            "phase5-compliance": {
                "mandatory": ["Maintain Quality Standards", "Regular Compliance Audits", "Monitoring Records"],
                "optional": ["Standards Manual", "Unauthorized Use Prevention", "Renewal Schedule"]
            }
        }

        # Document analysis: MoP qualitative review + optional document ensemble
        self.document_model = None
        self.document_feature_names = None
        self.explainer = None

        # Synonyms help match real IPOPHL uploads (authorization letters, receipts, etc.)
        self._term_synonyms = {
            "manual of specifications": [
                "manual of specification", "mop", "specifications manual", "code of practice",
                "technical specifications", "product specifications", "standard operating procedures",
                "sop", "spec manual"
            ],
            "causal link": [
                "causal relationship", "link between", "geographical area and", "connection to origin",
                "terroir", "link to locality", "specific to area", "origin-specific", "geographical link",
                "quality due to", "characteristics from area"
            ],
            "production process": [
                "production method", "processing method", "cultivation", "farming process",
                "growing practices", "harvesting", "post-harvest", "processing steps",
                "manufacturing process", "production flow"
            ],
            "quality control": [
                "quality standards", "quality assurance", "qc ", "quality check",
                "quality management", "qa", "quality inspection", "control measures",
                "quality verification"
            ],
            "labeling rules": [
                "labelling", "label requirements", "packaging rules", "labelling requirements",
                "packaging specifications", "marking rules", "label guidelines"
            ],
            "applicant entity": [
                "applicant", "organization", "association", "cooperative", "producers group",
                "farmers cooperative", "growers association", "registered organization",
                "legal entity", "applicant organization"
            ],
            "producers organization": [
                "producers", "farmers association", "growers", "membership", "producer group",
                "farmers group", "growers cooperative", "agricultural association"
            ],
            "legal standing": [
                "authorization", "authorized", "hereby authorize", "legal capacity", "representative",
                "legal authority", "authorized signatory", "power of attorney", "poa",
                "legal status", "registered entity"
            ],
            "membership list": [
                "members", "member list", "roster", "directory", "list of members",
                "membership roster", "member directory", "registered members"
            ],
            "stakeholder consultations": [
                "stakeholder", "consultation", "public hearing", "community consultation",
                "stakeholder meeting", "public consultation", "consultative meeting"
            ],
            "meeting minutes": [
                "minutes of meeting", "minutes of the", "meeting held", "minutes",
                "meeting notes", "meeting record", "minutes of proceedings"
            ],
            "consensus": [
                "agreed", "unanimous", "resolution", "agreement", "unanimous agreement",
                "collective decision", "general agreement"
            ],
            "governance board": [
                "board of directors", "governing board", "board resolution", "board members",
                "executive board", "advisory board", "board decisions"
            ],
            "technical validation": [
                "technical certification", "validated by", "certified by", "technical verification",
                "technical assessment", "certification", "validation", "technical approval"
            ],
            "government certification": [
                "department of agriculture", "bureau of", "da ", "certified", "government issued",
                "philippine government", "da certification", "bpi", "bureau of plant industry",
                "local government", "lgu certification", "city agriculturist", "municipal agriculturist"
            ],
            "independent verification": [
                "third party", "independent", "verified by", "third-party verification",
                "independent audit", "external verification", "independent assessment"
            ],
            "application form": [
                "application for", "duly accomplished", "application to register", "application",
                "form", "registration form", "application document", "filled-out form"
            ],
            "applicant name": [
                "name of applicant", "applicant's name", "applicant name", "entity name",
                "organization name", "applicant details"
            ],
            "domicile": [
                "address", "residence", "located at", "domiciled", "principal address",
                "registered address", "business address", "office address"
            ],
            "industrial establishment": [
                "establishment", "place of business", "office at", "business location",
                "facility", "plant", "production site", "processing plant"
            ],
            "file application": [
                "filed application", "filing", "submit application", "submitted",
                "application filed", "submission of application", "file registration"
            ],
            "bureau of trademarks": [
                "ipophl", "intellectual property", "bureau of trademark", "intellectual property office of the philippines",
                "trademark office", "ip office", "philippine ip office"
            ],
            "application package": [
                "application documents", "complete application", "submission package",
                "application dossier", "registration documents", "submission documents"
            ],
            "cover letter": [
                "letter of transmittal", "transmittal letter", "cover note", "letter of introduction",
                "accompanying letter"
            ],
            "official receipt": [
                "official receipt", "or number", "receipt no", "o.r.", "official receipt number",
                "payment receipt", "receipt"
            ],
            "application fee": [
                "filing fee", "payment of fee", "registration fee", "application fee",
                "processing fee", "fee payment"
            ],
            "proof of payment": [
                "proof of payment", "paid", "payment confirmation", "bank transfer",
                "payment slip", "deposit slip", "transaction receipt"
            ],
            "formality examination": [
                "formality", "formal examination", "formality check", "formal review",
                "examination of formalities"
            ],
            "substantive examination": [
                "substantive", "substance examination", "substantive review", "merits examination"
            ],
            "ip code compliance": [
                "intellectual property code", "ip code", "republic act", "ra 8293",
                "philippine ip code", "compliance with ip code"
            ],
            "deficiency notice response": [
                "deficiency", "response to notice", "comply with", "deficiency response",
                "response to deficiency", "correct deficiency", "remedy deficiency"
            ],
            "timeframe compliance": [
                "within the period", "deadline", "days from", "compliance with deadline",
                "on time", "timely submission", "within timeframe"
            ],
            "corrective actions": [
                "corrective action", "remedy", "amended", "correction", "corrective measures",
                "remedial actions", "amendment"
            ],
            "publication for opposition": [
                "publication", "opposition period", "published for opposition", "published",
                "publication in gazette", "ipophl gazette", "public notice"
            ],
            "public notice period": [
                "public notice", "notice period", "notice", "publication period"
            ],
            "opposition period": [
                "opposition", "third-party", "opposition phase", "opposition filing"
            ],
            "gi registration certificate": [
                "certificate of registration", "registration certificate", "gi certificate",
                "certificate", "gi registration", "registered gi certificate"
            ],
            "official notice of registration": [
                "notice of registration", "registered geographical indication", "registration notice",
                "official notice", "registration approval"
            ],
            "registration number": [
                "reg. no", "registration no", "certificate no", "registration number",
                "certificate number", "reg number"
            ],
            "maintain quality standards": [
                "quality standards", "maintain quality", "compliance with standards",
                "adhere to standards", "quality maintenance", "uphold standards"
            ],
            "regular compliance audits": [
                "audit", "compliance audit", "inspection", "regular audit", "compliance inspection",
                "monitoring audit"
            ],
            "monitoring records": [
                "monitoring", "records of", "documentation of", "monitoring logs",
                "records", "documentation", "monitoring documents"
            ],
            "lipa barako coffee": [
                "lipa coffee", "batangas coffee", "kapeng barako", "kape barako", "coffea liberica",
                "lipa liberica", "batangas barako", "lipa barako", "barako coffee"
            ],
            "product photos": [
                "photograph", "photographs", "product image", "product photo", "images",
                "photos", "product pictures"
            ],
            "roasting process": [
                "roasting process", "roasted coffee", "coffee roasting", "roast profile",
                "roasting method"
            ],
            "farming practices": [
                "farming practices", "farming practice", "traditional farming", "agricultural practices",
                "cultivation practices", "farm practices"
            ],
            "flavor profile": [
                "flavor profile", "taste profile", "cup profile", "taste notes", "flavor notes",
                "aroma profile", "sensory profile"
            ],
            "geographical origin": [
                "geographical origin", "geographical indication", "grown in lipa", "origin",
                "place of origin", "geographical source", "area of origin"
            ],
            "distinctive quality": [
                "distinctive quality", "unique quality", "characteristic quality", "unique characteristics",
                "distinct characteristics", "special quality"
            ],
            "geographical indication": [
                "geographical indication", " gi ", "indication of origin", "gi product",
                "geographically indicated"
            ],
            "gi": [
                "geographical indication", " gi registration", "protected gi", "geographical indication registration"
            ],
        }

        if ML_AVAILABLE and auto_train:
            self._initialize_models()
            self._ensure_document_model()
        elif ML_AVAILABLE:
            self._initialize_models()

    def _initialize_models(self):
        """Load GI document ensemble model (MoP advisory)."""
        doc_path = self.ml_dir / "gi_document_model.joblib"
        if doc_path.exists():
            try:
                self.document_model = joblib.load(doc_path)
                self.document_feature_names = (
                    ["text_length", "word_count"]
                    + self.gi_checklist["mandatory_terms"]
                    + self.gi_checklist["optional_terms"]
                )
                # SHAP explainer is created lazily on first explanation (slow to build)
                logging.info("Loaded document ML model from %s", doc_path.name)
            except Exception as e:
                logging.warning("Failed to load document ML model: %s", e)
                self.document_model = None

    def _ensure_official_mop_baseline(self) -> None:
        """Ensure official Part 1 / Part 2 / Control MoP CSV + training JSON exist."""
        try:
            try:
                from machinelearning.official_mop_dataset import (
                    DEFAULT_CSV_PATH,
                    DEFAULT_JSON_PATH,
                    build_official_mop_dataset,
                )
            except ImportError:
                from official_mop_dataset import (  # type: ignore[no-redef]
                    DEFAULT_CSV_PATH,
                    DEFAULT_JSON_PATH,
                    build_official_mop_dataset,
                )

            if not DEFAULT_CSV_PATH.exists() or not DEFAULT_JSON_PATH.exists():
                build_official_mop_dataset(self.ml_dir.parent)
                logging.info("Built official MoP dataset (Part 1, Part 2, Control & Traceability)")
        except Exception as exc:
            logging.warning("Could not build official MoP baseline dataset: %s", exc)

    def _ensure_document_model(self) -> None:
        """Train document classifier from official MoP files if missing."""
        self._ensure_official_mop_baseline()
        doc_path = self.ml_dir / "gi_document_model.joblib"
        if doc_path.exists() or self.document_model is not None:
            return
        try:
            import subprocess
            import sys

            script = self.ml_dir / "train_ai_model.py"
            if not script.exists():
                return
            proc = subprocess.run(
                [
                    sys.executable,
                    str(script),
                    "--train-documents",
                    "--data-dir",
                    str(self.ml_dir / "training_data"),
                ],
                capture_output=True,
                text=True,
                cwd=str(self.ml_dir.parent),
                timeout=300,
            )
            if proc.returncode == 0:
                self._initialize_models()
                logging.info("Auto-trained document RF from official MoP dataset (n1–n7)")
            else:
                logging.warning(
                    "Document model auto-train failed: %s",
                    (proc.stderr or proc.stdout or "")[-500:],
                )
        except Exception as e:
            logging.warning("Could not auto-train document model: %s", e)

    def _term_matches(self, text_lower: str, term: str) -> bool:
        """Match mandatory/optional GI terms with phrase, synonym, or multi-word co-occurrence."""
        term_lower = term.lower().strip()
        if not term_lower:
            return False
        if term_lower in text_lower:
            return True
        if " " not in term_lower:
            if re.search(r"\b" + re.escape(term_lower) + r"\b", text_lower):
                return True
        else:
            # All significant words present (e.g. "geographical" + "origin") without loose single-word synonyms
            parts = [p for p in re.split(r"\s+", term_lower) if len(p) > 3]
            if len(parts) >= 2 and all(p in text_lower for p in parts):
                return True
        for synonym in self._term_synonyms.get(term_lower, []):
            syn = synonym.lower().strip()
            if not syn:
                continue
            if syn in text_lower:
                return True
            if " " not in syn and re.search(r"\b" + re.escape(syn) + r"\b", text_lower):
                return True
        return False

    def _compute_readiness_score(
        self,
        detected_mandatory: List[str],
        missing_mandatory: List[str],
        detected_optional: List[str],
        checklist: Dict,
    ) -> int:
        """
        IPOPHL task score from explicit keyword coverage with term weighting:
        up to 70% from mandatory terms, up to 30% from optional terms.
        Terms with higher weights contribute more to the score.
        """
        mandatory_terms = checklist["mandatory_terms"]
        term_weights = checklist.get("term_weights", {})
        
        total_mandatory_weight = 0.0
        detected_mandatory_weight = 0.0
        for term in mandatory_terms:
            weight = term_weights.get(term, 1.0)
            total_mandatory_weight += weight
            if term in detected_mandatory:
                detected_mandatory_weight += weight
        
        mandatory_total = max(1.0, total_mandatory_weight)
        mandatory_score = (detected_mandatory_weight / mandatory_total) * 70
        
        optional_terms = checklist.get("optional_terms") or []
        optional_total = max(1, len(optional_terms)) if optional_terms else 1
        optional_score = (
            (len(detected_optional) / optional_total) * 30 if optional_terms else 0
        )
        
        return min(100, round(mandatory_score + optional_score))

    def _resolve_task_id_from_text(self, text: str, task_id: str | None) -> str | None:
        """Prefer upload-zone declared in document body (synthetic / test uploads)."""
        if not text:
            return task_id
        match = re.search(r"upload\s*zone:\s*(phase[1-5]-[\w-]+)", text, re.I)
        if not match:
            return task_id
        candidate = match.group(1).lower()
        if candidate in self.task_checklists:
            if task_id and task_id != candidate:
                logging.info(
                    "Resolved task_id from document content: %s (upload slot was %s)",
                    candidate,
                    task_id,
                )
            return candidate
        return task_id

    def ml_status(self) -> Dict:
        """Return whether the GI document model is available."""
        document_meta = {}
        document_path = self.ml_dir / "document_training_results.json"
        if document_path.exists():
            try:
                with open(document_path, encoding="utf-8") as f:
                    document_meta = json.load(f)
            except Exception:
                document_meta = {}
        return {
            "document_model_loaded": self.document_model is not None,
            "document_analysis_default": (
                "official_mop_ensemble_hybrid" if self.document_model else "official_mop_qualitative"
            ),
            "focus": "gi_document",
            "document_training": document_meta,
            "training": document_meta,
        }

    def _get_explainer(self):
        """Lazy-load SHAP TreeExplainer for the bagging (RF) member of the ensemble."""
        if self.explainer is not None:
            return self.explainer
        if not ML_AVAILABLE or self.document_model is None:
            return None
        try:
            try:
                from machinelearning.ensemble_learning import tree_estimator_for_shap
            except ImportError:
                from ensemble_learning import tree_estimator_for_shap  # type: ignore

            tree_model = tree_estimator_for_shap(self.document_model)
            self.explainer = shap.TreeExplainer(tree_model)
        except Exception as shap_err:
            logging.warning("SHAP explainer unavailable: %s", shap_err)
            self.explainer = None
        return self.explainer

    def _generate_shap_explanation(
        self,
        features: List,
        readiness_score: int,
        task_id: str = None,
        *,
        rule_result: Dict = None,
        rf_score: int = None,
        checklist: Dict = None,
    ) -> str:
        """SHAP narrative plus keyword checklist audit and ML/rule irregularities."""
        doc_type = task_id.replace("-", " ").title() if task_id else "Document"
        status = self._compliance_status_label(readiness_score)
        rule_score = int((rule_result or {}).get("readiness_score") or readiness_score)
        detected = (rule_result or {}).get("detected_features") or []
        missing = (rule_result or {}).get("missing_requirements") or []
        mandatory = (checklist or {}).get("mandatory_terms") or []
        optional = (checklist or {}).get("optional_terms") or []

        p1 = (
            f"<p>The keyword checklist for <strong>{doc_type}</strong> produced a readiness score of "
            f"<strong>{readiness_score}%</strong> ({status}). "
        )
        if mandatory:
            found_m = [t for t in mandatory if t in detected]
            p1 += (
                f"Mandatory terms matched: <strong>{len(found_m)} of {len(mandatory)}</strong>"
                f" ({', '.join(found_m) if found_m else 'none'}).</p>"
            )
        else:
            p1 += "</p>"

        irregularities: List[str] = []
        if rf_score is not None and abs(int(rf_score) - rule_score) >= 15:
            if int(rf_score) > rule_score:
                irregularities.append(
                    f"The ensemble model rated this document at <strong>{rf_score}%</strong>, "
                    f"which is <strong>{int(rf_score) - rule_score} points higher</strong> than the "
                    f"task keyword score ({rule_score}%). This suggests general GI vocabulary in the "
                    f"text may be inflating the statistical model even though required terms for this "
                    f"upload zone are still missing."
                )
            else:
                irregularities.append(
                    f"The ensemble model rated this document at <strong>{rf_score}%</strong>, "
                    f"which is <strong>{rule_score - int(rf_score)} points lower</strong> than the "
                    f"keyword score ({rule_score}%). The model may under-weight domain terms that "
                    f"are present but phrased differently from training samples."
                )

        shap_pos, shap_neg, _ = self._extract_shap_feature_impacts(features)
        for term in shap_pos:
            if term not in detected and term in mandatory + optional:
                irregularities.append(
                    f"SHAP flagged <strong>{term}</strong> as supportive, but that term was "
                    f"not matched by the strict task checklist for this upload zone."
                )

        p2 = "<p><strong>Gap analysis:</strong> "
        if missing:
            p2 += (
                f"Missing mandatory requirements: <strong>{', '.join(missing)}</strong>. "
            )
        else:
            p2 += "All mandatory keywords for this upload zone were found. "
        if optional:
            missing_opt = [t for t in optional if t not in detected]
            if missing_opt:
                p2 += (
                    f"Optional terms not detected: {', '.join(missing_opt[:4])}"
                    f"{'…' if len(missing_opt) > 4 else ''}. "
                )
        p2 += "</p>"

        if irregularities:
            p3 = (
                "<p><strong>ML irregularities:</strong> "
                + " ".join(irregularities)
                + "</p>"
            )
        elif shap_pos or shap_neg:
            p3 = "<p><strong>Model feature influence:</strong> "
            if shap_pos:
                p3 += f"Terms that increased ML confidence: {', '.join(shap_pos)}. "
            if shap_neg:
                p3 += f"Terms that decreased ML confidence: {', '.join(shap_neg)}. "
            p3 += "The final score follows the task keyword checklist, not the ML model alone.</p>"
        else:
            p3 = (
                "<p>The keyword score and ML signals are aligned for this document. "
                "Re-upload after adding any missing mandatory terms to improve readiness.</p>"
            )

        return p1 + p2 + p3

    def _extract_shap_feature_impacts(
        self, features: List
    ) -> tuple[List[str], List[str], List[dict]]:
        """Return (positive_terms, negative_terms, full_impact_rows) from TreeExplainer."""
        feature_names = self.document_feature_names or []
        explainer = self._get_explainer()
        if not explainer or not feature_names:
            return [], [], []

        try:
            shap_values = explainer.shap_values(np.array([features]))
            if isinstance(shap_values, list):
                instance_shap = shap_values[1][0]
            else:
                instance_shap = (
                    shap_values[0, :, 1] if len(shap_values.shape) == 3 else shap_values[0]
                )
            feature_impact: List[dict] = []
            for i, val in enumerate(instance_shap):
                if i >= len(feature_names):
                    break
                name = feature_names[i]
                if name in ("text_length", "word_count"):
                    continue
                feature_impact.append({"name": name, "impact": float(val)})
            feature_impact.sort(key=lambda x: abs(x["impact"]), reverse=True)
            shap_pos = [f["name"] for f in feature_impact if f["impact"] > 0.05][:5]
            shap_neg = [f["name"] for f in feature_impact if f["impact"] < -0.05][:5]
            return shap_pos, shap_neg, feature_impact[:8]
        except Exception as exc:
            logging.warning("SHAP value extraction failed: %s", exc)
            return [], [], []

    def _build_rf_shap_appendix(
        self,
        features: List,
        rf_score: int,
        mop_status: str,
        task_id: str | None = None,
    ) -> str:
        """Append ensemble probability + SHAP term impacts for capstone validation (advisory)."""
        doc_type = task_id.replace("-", " ").title() if task_id else "Document"
        rf_ready = int(rf_score) >= 75
        mop_ready = str(mop_status).lower() == "ready"
        agrees = rf_ready == mop_ready

        shap_pos, shap_neg, impacts = self._extract_shap_feature_impacts(features)
        agreement = (
            "<strong>aligned</strong> — MoP review and the ensemble agree."
            if agrees
            else "<strong>divergent</strong> — treat MoP Ready/Not Ready as final; "
            "review ensemble signals below."
        )

        parts = [
            f"<p><strong>Ensemble validation ({doc_type}):</strong> "
            f"soft-voting confidence <strong>{int(rf_score)}%</strong> "
            f"({'Ready' if rf_ready else 'Not Ready'} at 75% threshold). "
            f"MoP classification: <strong>{mop_status}</strong> — {agreement} "
            f"Base learners: Random Forest + Extra Trees (bagging) and Gradient Boosting "
            f"(boosting).</p>"
        ]

        if impacts:
            rows = ", ".join(
                f"{row['name']} ({row['impact']:+.3f})" for row in impacts[:6]
            )
            parts.append(
                f"<p><strong>SHAP feature influence (RF bagging member):</strong> {rows}. "
                "Positive values push toward Ready; negative values push toward Not Ready.</p>"
            )
        elif shap_pos or shap_neg:
            detail = []
            if shap_pos:
                detail.append(f"supportive: {', '.join(shap_pos)}")
            if shap_neg:
                detail.append(f"against Ready: {', '.join(shap_neg)}")
            parts.append(
                f"<p><strong>SHAP feature influence:</strong> {'; '.join(detail)}.</p>"
            )

        if not agrees:
            if mop_ready and not rf_ready:
                parts.append(
                    "<p>The MoP themes are satisfied, but the statistical ensemble is cautious — "
                    "consider adding more explicit Kapeng Barako / Lipa terminology to strengthen "
                    "ML agreement.</p>"
                )
            else:
                parts.append(
                    "<p>The ensemble may reflect GI vocabulary in the text, but MoP critical "
                    "themes are still missing or the document targets the wrong product — follow "
                    "the gap analysis above.</p>"
                )

        return "".join(parts)

    @staticmethod
    def _compliance_status_label(score: int) -> str:
        if score >= 85:
            return "highly compliant"
        if score >= 70:
            return "conditionally sufficient"
        return "insufficient"

    def _align_shap_readiness_text(self, html: str, score: int) -> str:
        """Ensure narrative text uses the same readiness % as the score bar."""
        if not html:
            return html
        score = max(0, min(100, int(score)))
        status_label = self._compliance_status_label(score)
        out = re.sub(
            r"(readiness score of\s*<strong>)\d+(%</strong>)",
            rf"\g<1>{score}\g<2>",
            html,
            flags=re.IGNORECASE,
        )
        out = re.sub(
            r"(initial readiness score of\s*<strong>)\d+(%</strong>)",
            rf"\g<1>{score}\g<2>",
            out,
            flags=re.IGNORECASE,
        )
        out = re.sub(
            r"status of\s*<strong>[^<]*</strong>",
            f"status of <strong>{status_label}</strong>",
            out,
            count=1,
            flags=re.IGNORECASE,
        )
        return out

    def normalize_analysis_payload(self, payload: Dict) -> Dict:
        """Single canonical readiness_score across UI fields and SHAP narrative."""
        if not payload:
            return payload
        normalized = dict(payload)
        score = int(normalized.get("readiness_score") or 0)
        shap = normalized.get("shap_analysis") or ""
        if shap:
            normalized["shap_analysis"] = self._align_shap_readiness_text(shap, score)
        return normalized

    def _normalize_extracted_text(self, text: str) -> str:
        """Clean OCR/PDF noise and collapse whitespace for reliable term matching."""
        cleaned = (text or "").replace("\r\n", "\n").replace("\r", "\n")
        cleaned = re.sub(r"[ \t\f\v]+", " ", cleaned)
        cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
        return cleaned.strip()

    def _analysis_scan_meta(self, text: str) -> Dict:
        words = re.findall(r"[A-Za-z0-9']+", text or "")
        excerpt = (text or "").strip()[:600]
        digest = hashlib.sha256((text or "").encode("utf-8", errors="ignore")).hexdigest()
        return {
            "word_count": len(words),
            "text_excerpt": excerpt,
            "content_fingerprint": digest[:16],
        }

    def extract_text_from_file(self, file_path: str) -> str:
        """Extract text from uploaded file"""
        file_path = Path(file_path)

        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        suffix = file_path.suffix.lower()
        text = ""

        if suffix == '.pdf':
            text = self._extract_from_pdf(file_path)
        elif suffix in ['.doc', '.docx']:
            if suffix == '.doc':
                raise ValueError(
                    "Legacy .doc files are not supported for text scanning. "
                    "Re-save as .docx or export to PDF, then upload again."
                )
            text = self._extract_from_docx(file_path)
        elif suffix in ['.txt', '.md']:
            text = file_path.read_text(encoding='utf-8', errors='ignore')
        else:
            raise ValueError(f"Unsupported file type: {file_path.suffix}")

        return self._normalize_extracted_text(text)

    def _extract_from_pdf(self, file_path: Path) -> str:
        """Extract text from PDF, with OCR fallback for scanned pages."""
        if not PDF_AVAILABLE:
            raise ImportError("PyMuPDF is required for PDF processing")

        chunks: list[str] = []
        try:
            doc = fitz.open(file_path)
            for page in doc:
                page_text = page.get_text("text", sort=True) or page.get_text() or ""
                page_text = page_text.strip()
                if len(page_text) < 40 and OCR_AVAILABLE:
                    page_text = self._ocr_pdf_page(page) or page_text
                if page_text:
                    chunks.append(page_text)
            doc.close()
        except Exception as e:
            logging.error(f"Error extracting from PDF: {e}")
            if OCR_AVAILABLE:
                return self._ocr_pdf(file_path)

        text = "\n\n".join(chunks)
        if len(text.strip()) < 80 and OCR_AVAILABLE:
            ocr_text = self._ocr_pdf(file_path)
            if len(ocr_text.strip()) > len(text.strip()):
                text = ocr_text
        return text

    def _ocr_pdf_page(self, page) -> str:
        if not OCR_AVAILABLE:
            return ""
        try:
            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
            img_data = pix.tobytes("png")
            img = Image.open(io.BytesIO(img_data))
            return pytesseract.image_to_string(img)
        except Exception as e:
            logging.error(f"OCR page error: {e}")
            return ""

    def _ocr_pdf(self, file_path: Path) -> str:
        """OCR extraction from PDF pages"""
        if not OCR_AVAILABLE:
            return ""

        text = ""
        try:
            doc = fitz.open(file_path)
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                pix = page.get_pixmap()
                img_data = pix.tobytes("png")
                img = Image.open(io.BytesIO(img_data))
                page_text = pytesseract.image_to_string(img)
                text += page_text + "\n"
            doc.close()
        except Exception as e:
            logging.error(f"OCR error: {e}")

        return text

    def _extract_from_docx(self, file_path: Path) -> str:
        """Extract text from Word document including tables, headers, and footers."""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required for Word document processing")

        parts: list[str] = []
        try:
            doc = docx.Document(file_path)
            for paragraph in doc.paragraphs:
                line = paragraph.text.strip()
                if line:
                    parts.append(line)

            for table in doc.tables:
                for row in table.rows:
                    row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_cells:
                        parts.append(" | ".join(row_cells))

            for section in doc.sections:
                for container in (section.header, section.footer):
                    if container is None:
                        continue
                    for paragraph in container.paragraphs:
                        line = paragraph.text.strip()
                        if line:
                            parts.append(line)
                    for table in container.tables:
                        for row in table.rows:
                            row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                            if row_cells:
                                parts.append(" | ".join(row_cells))
        except Exception as e:
            logging.error(f"Error extracting from DOCX: {e}")
            raise ValueError(f"Could not read DOCX content: {e}") from e

        return "\n".join(parts)

    def _attach_scan_meta(self, payload: Dict, text: str) -> Dict:
        meta = self._analysis_scan_meta(text)
        enriched = dict(payload or {})
        enriched["text_excerpt"] = meta["text_excerpt"]
        enriched["content_fingerprint"] = meta["content_fingerprint"]
        enriched["word_count"] = meta["word_count"]
        breakdown = enriched.get("score_breakdown")
        if isinstance(breakdown, dict):
            merged_breakdown = dict(breakdown)
            merged_breakdown.update(meta)
            enriched["score_breakdown"] = merged_breakdown
        insights = enriched.get("ip_pillar_assessment")
        if isinstance(insights, dict):
            doc_insights = dict(insights.get("document_insights") or {})
            doc_insights.setdefault("word_count", meta["word_count"])
            doc_insights["text_excerpt"] = meta["text_excerpt"]
            doc_insights["content_fingerprint"] = meta["content_fingerprint"]
            insights = dict(insights)
            insights["document_insights"] = doc_insights
            enriched["ip_pillar_assessment"] = insights
        return enriched

    def analyze_document(self, file_path: str, task_id: str = None) -> Dict:
        """Main analysis function with task-specific context"""
        file_path_obj = Path(file_path)

        if file_path_obj.suffix.lower() == ".csv":
            return {
                "success": False,
                "error": "CSV farmer-profile analysis was removed. Upload a GI document (PDF/DOCX/TXT) for MoP analysis.",
            }

        try:
            # Extract text
            text = self.extract_text_from_file(file_path)
            scan_meta = self._analysis_scan_meta(text)
            if scan_meta["word_count"] < 40:
                return {
                    "success": False,
                    "error": (
                        "Could not extract enough readable text from this file. "
                        "Try a text-based PDF or DOCX, or re-scan the document with higher quality."
                    ),
                    "readiness_score": 0,
                    "status": "Not Ready",
                    "detected_features": [],
                    "missing_requirements": ["Readable document text"],
                    "text_length": len(text or ""),
                    "text_excerpt": scan_meta.get("text_excerpt") or "",
                    "content_fingerprint": scan_meta.get("content_fingerprint") or "",
                    "word_count": scan_meta.get("word_count") or 0,
                }

            task_id = self._resolve_task_id_from_text(text, task_id)

            # Default IPOPHL pipeline: official MoP qualitative review + RF/SHAP validation
            self._ensure_document_model()

            # Determine which checklist to use
            checklist = self.gi_checklist
            if task_id in self.task_checklists:
                checklist = {
                    "mandatory_terms": self.task_checklists[task_id]["mandatory"],
                    "optional_terms": self.task_checklists[task_id]["optional"]
                }
                logging.info(f"Using task-specific checklist for: {task_id}")

            rule_result = self._rule_based_analysis(text, checklist, task_id=task_id)
            if self.document_model is not None:
                ml_result = self._ml_analysis(text, checklist, task_id)
                merged = self._merge_analysis_results(
                    rule_result, ml_result, text=text, task_id=task_id, checklist=checklist
                )
                if task_id:
                    merged["task_id"] = task_id
                return self._attach_scan_meta(merged, text)
            payload = self.normalize_analysis_payload(rule_result)
            payload["analysis_method"] = "official_mop_qualitative"
            if task_id:
                payload["task_id"] = task_id
            return self._attach_scan_meta(payload, text)

        except Exception as e:
            logging.error(f"Analysis error: {e}")
            
            # Fallback checklist
            checklist = self.gi_checklist
            if task_id in self.task_checklists:
                checklist = {
                    "mandatory_terms": self.task_checklists[task_id]["mandatory"],
                    "optional_terms": self.task_checklists[task_id]["optional"]
                }
                
            return {
                "success": False,
                "error": str(e),
                "readiness_score": 0,
                "status": "Not Ready",
                "detected_features": [],
                "missing_requirements": checklist["mandatory_terms"]
            }

    def _rule_based_analysis(self, text: str, checklist: Dict = None, task_id: str | None = None) -> Dict:
        """Qualitative MoP-basis analysis (PART 1 / PART 2 / Control & Labelling).

        Does not use a keyword percentage readiness score. Status is Ready / Not Ready
        from theme coverage against the Kapeng Barako specification drafting package.
        """
        if checklist is None:
            checklist = self.gi_checklist

        try:
            from machinelearning.gi_reference_basis import evaluate_against_reference
            from machinelearning.ip_pillars import evaluate_ip_pillars
        except ImportError:
            from gi_reference_basis import evaluate_against_reference
            from ip_pillars import evaluate_ip_pillars

        text_lower = (text or "").lower()
        review = evaluate_against_reference(
            text or "",
            task_id=task_id,
            term_matches=self._term_matches,
        )
        status = review.get("status") or "Not Ready"
        # Persist 100/0 for legacy DB fields only — never shown as a readiness %.
        readiness_score = 100 if status == "Ready" else 0
        strengths = list(review.get("strengths") or [])
        missing = list(review.get("missing_requirements") or [])
        detected = list(review.get("detected_features") or strengths)
        themes = list(review.get("themes") or [])
        improvements = list(review.get("improvements") or [])

        sections = [
            {
                "id": t.get("id"),
                "label": t.get("label"),
                "found": t.get("coverage") == "well_covered",
                "coverage": t.get("coverage"),
                "part": t.get("part"),
                "expectation": t.get("expectation"),
                "evidence": t.get("evidence_signals") or [],
            }
            for t in themes
        ]
        score_breakdown = {
            "analysis_mode": "mop_reference_qualitative",
            "reference_source": review.get("reference_source"),
            "final_score": readiness_score,
            "formula": "Qualitative GI theme coverage (no keyword percentage)",
            "sections": sections,
            "sections_found": sum(1 for s in sections if s.get("found")),
            "sections_total": len(sections),
            "improvements": improvements,
            "product_focus": review.get("product_focus"),
            "terms": [
                {
                    "term": t.get("label"),
                    "found": t.get("coverage") != "missing",
                    "coverage": t.get("coverage"),
                }
                for t in themes
            ],
        }

        ip_pillar_assessment = evaluate_ip_pillars(
            text_lower,
            detected_features=detected,
            missing_requirements=missing,
            rubric_sections=sections,
            term_matches=self._term_matches,
            task_id=task_id,
            document_ready=(status == "Ready" and bool((review.get("product_focus") or {}).get("ok", True))),
            source_text=text,
            text_length=len(text or ""),
            mandatory_met=sum(1 for t in themes if t.get("critical") and t.get("coverage") != "missing"),
            mandatory_total=max(1, sum(1 for t in themes if t.get("critical"))),
        )
        if isinstance(ip_pillar_assessment, dict):
            ip_pillar_assessment["executive_summary"] = re.sub(
                r"<[^>]+>",
                " ",
                (review.get("shap_analysis") or "").split("</p>")[0],
            ).strip()
            ip_pillar_assessment["recommendations"] = improvements
            ip_pillar_assessment["document_insights"] = {
                "document_type": (task_id or "gi-document").replace("-", " ").title(),
                "word_count": review.get("word_count") or 0,
                "checklist_met": sum(1 for t in themes if t.get("coverage") == "well_covered"),
                "checklist_total": len(themes),
                "detected_features": strengths[:10],
                "missing_requirements": missing[:10],
                "reference_source": review.get("reference_source"),
            }

        return {
            "success": True,
            "readiness_score": readiness_score,
            "status": status,
            "detected_features": detected[:40],
            "missing_requirements": missing,
            "text_length": len(text or ""),
            "analysis_method": "mop_reference_qualitative",
            "shap_analysis": review.get("shap_analysis") or "",
            "score_breakdown": score_breakdown,
            "keyword_score": None,
            "section_score": None,
            "ip_pillar_assessment": ip_pillar_assessment,
            "improvements": improvements,
            "product_focus": review.get("product_focus"),
        }

    def _task_keyword_ml_score(self, text: str, checklist: Dict) -> int:
        """ML display score mirrors task keyword coverage (same formula as rules)."""
        text_lower = text.lower()
        detected_m = [
            t for t in checklist["mandatory_terms"] if self._term_matches(text_lower, t)
        ]
        missing_m = [t for t in checklist["mandatory_terms"] if t not in detected_m]
        detected_o = [
            t for t in checklist.get("optional_terms", []) if self._term_matches(text_lower, t)
        ]
        return self._compute_readiness_score(detected_m, missing_m, detected_o, checklist)

    def _random_forest_document_score(self, features: List) -> int | None:
        if self.document_model is None:
            return None
        try:
            if hasattr(self.document_model, "predict_proba"):
                probability = self.document_model.predict_proba([features])[0]
                ready_idx = 1 if len(probability) > 1 else 0
                return int(round(float(probability[ready_idx]) * 100))
            return int(self.document_model.predict([features])[0] * 100)
        except Exception:
            return None

    def _merge_analysis_results(
        self,
        rule_result: Dict,
        ml_result: Dict,
        *,
        text: str = None,
        task_id: str = None,
        checklist: Dict = None,
    ) -> Dict:
        """Qualitative MoP review is authoritative; RF score is advisory only."""
        rule_score = int(rule_result.get("readiness_score") or 0)
        ml_score = int(ml_result.get("readiness_score") or 0)
        rf_score = ml_result.get("rf_score")
        merged_score = rule_score
        detected = list(rule_result.get("detected_features") or [])
        missing = list(rule_result.get("missing_requirements") or [])
        status = rule_result.get("status") or ("Ready" if merged_score >= 100 else "Not Ready")

        shap = rule_result.get("shap_analysis") or ""
        if not shap:
            shap = self._keyword_shap_fallback(rule_result, merged_score, task_id)

        rf_agreement = None
        if rf_score is not None and text:
            rf_ready = int(rf_score) >= 75
            mop_ready = str(status).lower() == "ready"
            rf_agreement = rf_ready == mop_ready

        return self.normalize_analysis_payload({
            "success": True,
            "readiness_score": merged_score,
            "status": status,
            "detected_features": detected,
            "missing_requirements": missing,
            "text_length": rule_result.get("text_length") or ml_result.get("text_length") or 0,
            "analysis_method": "official_mop_ensemble_hybrid",
            "shap_analysis": shap,
            "ml_score": ml_score,
            "rule_score": rule_score,
            "rf_score": rf_score,
            "rf_agreement": rf_agreement,
            "score_breakdown": {
                **(rule_result.get("score_breakdown") or {}),
                "rf_validation": {
                    "rf_score": rf_score,
                    "rf_agreement": rf_agreement,
                    "training_source": "ipophl_official_mop_dataset.csv",
                    "ensemble_method": "soft_voting_bagging_boosting",
                },
            },
            "keyword_score": None,
            "section_score": None,
            "ip_pillar_assessment": rule_result.get("ip_pillar_assessment"),
            "improvements": rule_result.get("improvements"),
        })

    def _keyword_shap_fallback(self, rule_result: Dict, score: int, task_id: str | None) -> str:
        doc_type = task_id.replace("-", " ").title() if task_id else "Document"
        detected = rule_result.get("detected_features") or []
        missing = rule_result.get("missing_requirements") or []
        status = rule_result.get("status") or ("Ready" if score >= 100 else "Not Ready")
        p1 = (
            f"<p>Qualitative AI review for <strong>{doc_type}</strong>: "
            f"classification <strong>{status}</strong>. "
            f"Strengths detected: {', '.join(detected[:8]) if detected else 'none yet'}.</p>"
        )
        p2 = (
            f"<p>Themes still needing work: {', '.join(missing) if missing else 'none'}. "
            f"Revise against the Kapeng Barako GI filing requirements, then re-analyze.</p>"
        )
        return p1 + p2

    def _ml_analysis(self, text: str, checklist: Dict = None, task_id: str = None) -> Dict:
        """ML layer: task keyword score + ensemble advisory probability."""
        if checklist is None:
            checklist = self.gi_checklist

        try:
            features = self._extract_features(text)
            rf_score = self._random_forest_document_score(features)
            readiness_score = self._task_keyword_ml_score(text, checklist)
            detected_features, missing_requirements = self._analyze_terms(text, checklist)
            shap_analysis = self._generate_shap_explanation(
                features,
                readiness_score,
                task_id,
                rule_result={
                    "readiness_score": readiness_score,
                    "detected_features": detected_features,
                    "missing_requirements": missing_requirements,
                },
                rf_score=rf_score,
                checklist=checklist,
            )

            return self.normalize_analysis_payload({
                "success": True,
                "readiness_score": readiness_score,
                "status": "Ready" if readiness_score >= 75 else "Not Ready",
                "detected_features": detected_features,
                "missing_requirements": missing_requirements,
                "text_length": len(text),
                "analysis_method": "ml_based",
                "shap_analysis": shap_analysis,
                "rf_score": rf_score,
                "ml_score": readiness_score,
            })

        except Exception as e:
            logging.error(f"ML analysis failed, falling back to rule-based: {e}")
            return self.normalize_analysis_payload(self._rule_based_analysis(text, checklist))

    def _extract_features(self, text: str) -> List:
        """Extract ML features from text (fixed global vector for trained RF model)."""
        features = [len(text), len(text.split())]
        all_terms = self.gi_checklist["mandatory_terms"] + self.gi_checklist["optional_terms"]
        text_lower = text.lower()
        for term in all_terms:
            features.append(1 if self._term_matches(text_lower, term) else 0)
        return features

    def _analyze_terms(self, text: str, checklist: Dict = None) -> Tuple[List[str], List[str]]:
        """Analyze which terms are present/missing from checklist"""
        if checklist is None:
            checklist = self.gi_checklist
            
        text_lower = text.lower()

        detected = []
        missing = []

        for term in checklist["mandatory_terms"]:
            if self._term_matches(text_lower, term):
                detected.append(term)
            else:
                missing.append(term)
        
        # Also check optional terms for "detected"
        for term in checklist["optional_terms"]:
            if self._term_matches(text_lower, term):
                detected.append(term)

        return detected, missing

    def save_uploaded_file(self, file_data, filename: str) -> str:
        """Save uploaded file with UUID naming"""
        # Generate unique filename
        file_uuid = str(uuid.uuid4())
        file_ext = Path(filename).suffix
        safe_filename = f"{file_uuid}{file_ext}"

        # Save file
        file_path = self.uploads_dir / safe_filename
        if hasattr(file_data, 'save'):
            file_data.save(str(file_path))
        else:
            with open(file_path, 'wb') as f:
                f.write(file_data)

        return str(file_path)

    def get_file_preview_url(self, file_path: str) -> str:
        """Get URL for file preview"""
        return f"/api/file-preview/{Path(file_path).name}"

# Global instance (lazy train — avoids racing dataset build on import)
gi_analyzer = GIAnalyzer(auto_train=False)
