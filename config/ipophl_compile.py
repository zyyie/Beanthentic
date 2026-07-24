"""
Compile IPOPHL Phase 1–3 uploads into a single PDF or DOCX package.

DOCX: deep-copies source body paragraphs/tables in order (tables unchanged).
PDF: uses a Unicode system font so Filipino / smart punctuation is not replaced with '?'.
"""

from __future__ import annotations

import copy
import html
import io
import re
from datetime import datetime
from pathlib import Path

from config.ipophl_store import (
    OFFICIAL_IPOPHL_TASK_IDS,
    get_document,
    list_documents,
    resolve_file_path,
    task_label,
)

_IMAGE_SUFFIXES = frozenset({".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp"})
_TEXT_SUFFIXES = frozenset({".txt", ".md", ".csv"})

# Windows / common Unicode fonts (Helvetica cannot render many Word characters).
_UNICODE_FONT_CANDIDATES = (
    Path(r"C:\Windows\Fonts\arial.ttf"),
    Path(r"C:\Windows\Fonts\calibri.ttf"),
    Path(r"C:\Windows\Fonts\segoeui.ttf"),
    Path(r"C:\Windows\Fonts\times.ttf"),
    Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
    Path("/System/Library/Fonts/Supplemental/Arial.ttf"),
)


def collect_compile_sources(
    *,
    file_uuids: list[str] | None = None,
) -> list[dict]:
    """
    Ordered list of source files for the compiled package.
    Prefer official MoP zones, then any extra UUID list.
    """
    sources: list[dict] = []
    seen: set[str] = set()

    def add_record(record: dict | None) -> None:
        if not record:
            return
        uid = str(record.get("file_uuid") or "").strip()
        if not uid or uid in seen:
            return
        hint = str(record.get("original_filename") or "")
        path = resolve_file_path(uid, filename_hint=hint or None)
        if not path or not path.is_file() or path.stat().st_size <= 0:
            return
        tid = str(record.get("task_id") or "ipophl-other")
        seen.add(uid)
        sources.append(
            {
                "file_uuid": uid,
                "task_id": tid,
                "label": task_label(tid),
                "original_filename": hint or path.name,
                "path": path,
            }
        )

    if file_uuids:
        for raw in file_uuids:
            uid = str(raw or "").strip()
            if uid:
                add_record(get_document(uid))
        if sources:
            return sources

    for tid in OFFICIAL_IPOPHL_TASK_IDS:
        docs = list_documents(task_id=tid, limit=50)
        docs_sorted = sorted(
            docs,
            key=lambda d: str(d.get("upload_timestamp") or d.get("created_at") or ""),
            reverse=True,
        )
        for record in docs_sorted:
            add_record(record)

    if not sources:
        for record in list_documents(limit=300):
            tid = str(record.get("task_id") or "")
            if tid.startswith(("phase1-", "phase2-", "phase3-")):
                add_record(record)

    return sources


def _clean_text(value: str) -> str:
    """Normalize Word/XML quirks without inventing '?' placeholders."""
    if not value:
        return ""
    text = str(value)
    # Remove U+FFFD replacement chars and NUL; keep real punctuation.
    text = text.replace("\ufffd", "").replace("\x00", "")
    # Normalize common Word whitespace
    text = text.replace("\xa0", " ").replace("\u200b", "")
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    return text


def _resolve_unicode_font() -> Path | None:
    for path in _UNICODE_FONT_CANDIDATES:
        if path.is_file():
            return path
    return None


def _extract_text_from_docx(path: Path) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("python-docx is required to compile Word documents.") from exc

    doc = Document(str(path))
    parts: list[str] = []
    for block in _iter_docx_blocks(doc):
        kind, obj = block
        if kind == "p":
            text = _clean_text(obj.text or "").strip()
            if text:
                parts.append(text)
        elif kind == "tbl":
            # Preserve table rows as tab-separated lines for PDF/text fallback only.
            for row in obj.rows:
                cells = [_clean_text(c.text or "").strip() for c in row.cells]
                # Deduplicate merged-cell repeats from python-docx
                line = "\t".join(cells)
                if line.strip():
                    parts.append(line)
    return "\n\n".join(parts).strip()


def _extract_text_from_pdf(path: Path) -> str:
    import fitz

    parts: list[str] = []
    with fitz.open(str(path)) as pdf:
        for page in pdf:
            parts.append(_clean_text(page.get_text("text") or "").strip())
    return "\n\n".join(p for p in parts if p).strip()


def _extract_plain_text(path: Path) -> str:
    raw = path.read_bytes()
    for enc in ("utf-8", "utf-8-sig", "cp1252", "latin-1"):
        try:
            return _clean_text(raw.decode(enc)).strip()
        except UnicodeDecodeError:
            continue
    return _clean_text(raw.decode("utf-8", errors="ignore")).strip()


def _source_body_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return _extract_text_from_docx(path)
    if suffix == ".pdf":
        return _extract_text_from_pdf(path)
    if suffix in _TEXT_SUFFIXES:
        return _extract_plain_text(path)
    if suffix == ".doc":
        return (
            f"[Legacy .doc file: {path.name}]\n"
            "Re-save as .docx or PDF for full text inclusion in the compiled package."
        )
    return f"[Embedded file: {path.name}]"


def _iter_docx_blocks(document):
    """Yield ('p', paragraph) / ('tbl', table) in true document-body order."""
    from docx.oxml.ns import qn
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    body = document.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield ("p", Paragraph(child, document))
        elif child.tag == qn("w:tbl"):
            yield ("tbl", Table(child, document))


def _append_page_break(doc) -> None:
    from docx.enum.text import WD_BREAK

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def _copy_docx_source_into(doc, src_path: Path) -> int:
    """
    Deep-copy paragraphs and tables from src into doc in body order.
    Tables are cloned as XML (structure/content preserved, not rebuilt as text).
    """
    from docx import Document as SrcDocument
    from docx.oxml.ns import qn

    src = SrcDocument(str(src_path))
    copied = 0
    dest_body = doc.element.body

    # Insert before sectPr if present
    sect = dest_body.find(qn("w:sectPr"))

    for kind, obj in _iter_docx_blocks(src):
        if kind == "p":
            text = _clean_text(obj.text or "")
            # Keep intentional blank paragraphs sparingly; skip empty runs noise
            if not text.strip():
                # Still copy paragraph if it only hosts a page break / drawing? skip empty
                continue
            element = copy.deepcopy(obj._element)
            if sect is not None:
                sect.addprevious(element)
            else:
                dest_body.append(element)
            copied += 1
        elif kind == "tbl":
            element = copy.deepcopy(obj._element)
            if sect is not None:
                sect.addprevious(element)
            else:
                dest_body.append(element)
            copied += 1

    return copied


def build_compiled_docx(sources: list[dict]) -> bytes:
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt, RGBColor
    except ImportError as exc:
        raise RuntimeError("python-docx is required to build a DOCX package.") from exc

    doc = Document()
    title = doc.add_heading("Kapeng Barako - Compiled GI Documentation", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(
        f"Beanthentic IPOPHL compile - {datetime.now().strftime('%Y-%m-%d %H:%M')}\n"
        f"{len(sources)} source file(s) from Phases 1-3"
    )
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph(
        "This package combines Justification, Technical, and Control & Traceability "
        "documents into one file for review and submission preparation. "
        "Source tables are included unchanged."
    )

    for idx, source in enumerate(sources):
        if idx > 0:
            _append_page_break(doc)

        path: Path = source["path"]
        label = str(source.get("label") or "Document")
        fname = str(source.get("original_filename") or path.name)
        task_id = str(source.get("task_id") or "")

        doc.add_heading(_clean_text(label), level=1)
        info = doc.add_paragraph()
        info_run = info.add_run(f"Source: {_clean_text(fname)}")
        if task_id:
            info.add_run(f"  |  Zone: {task_id}")
        info_run.font.size = Pt(9)
        info_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        suffix = path.suffix.lower()
        if suffix in _IMAGE_SUFFIXES:
            try:
                doc.add_picture(str(path), width=Inches(6.0))
            except Exception:
                doc.add_paragraph(f"[Could not embed image: {fname}]")
            continue

        if suffix == ".docx":
            try:
                if _copy_docx_source_into(doc, path) > 0:
                    continue
            except Exception:
                pass

        body = _source_body_text(path)
        if not body:
            doc.add_paragraph(f"[No extractable text from {fname}]")
            continue
        for block in body.split("\n\n"):
            line = _clean_text(block).strip()
            if line:
                doc.add_paragraph(line)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _pdf_insert_text(page, rect, text: str, *, fontsize: float, fontfile: Path | None, color=(0.1, 0.1, 0.1), align=0) -> None:
    """Insert Unicode text; prefer embedded TTF so characters are not replaced with '?'."""
    import fitz

    clean = _clean_text(text)
    if not clean:
        return
    if fontfile and fontfile.is_file():
        page.insert_textbox(
            rect,
            clean,
            fontsize=fontsize,
            fontfile=str(fontfile),
            fontname="unicode",
            color=color,
            align=align,
        )
    else:
        # Last resort: map unsupported glyphs to closest ASCII (never leave U+FFFD as '?')
        ascii_ish = (
            clean.replace("—", "-")
            .replace("–", "-")
            .replace("•", "*")
            .replace("·", "|")
            .replace("“", '"')
            .replace("”", '"')
            .replace("‘", "'")
            .replace("’", "'")
            .replace("…", "...")
        )
        ascii_ish = re.sub(r"[^\x09\x0a\x0d\x20-\x7e]", "", ascii_ish)
        page.insert_textbox(
            rect,
            ascii_ish,
            fontsize=fontsize,
            fontname="helv",
            color=color,
            align=align,
        )


def _table_to_html(table) -> str:
    rows_html: list[str] = []
    seen_row_keys: set[str] = set()
    for row in table.rows:
        cells = [_clean_text(c.text or "").strip() for c in row.cells]
        key = "||".join(cells)
        if key in seen_row_keys:
            continue
        seen_row_keys.add(key)
        tds = "".join(
            f"<td style='border:1px solid #444;padding:4px 6px;vertical-align:top;'>{html.escape(c)}</td>"
            for c in cells
        )
        rows_html.append(f"<tr>{tds}</tr>")
    if not rows_html:
        return ""
    return (
        "<table style='border-collapse:collapse;width:100%;font-size:10px;margin:8px 0;'>"
        + "".join(rows_html)
        + "</table>"
    )


def _docx_to_html_chunks(path: Path) -> list[str]:
    """Build HTML chunks (paragraphs + unchanged table markup) for PDF rendering."""
    from docx import Document

    doc = Document(str(path))
    chunks: list[str] = []
    buf: list[str] = []

    def flush() -> None:
        nonlocal buf
        if buf:
            chunks.append("".join(buf))
            buf = []

    for kind, obj in _iter_docx_blocks(doc):
        if kind == "p":
            text = _clean_text(obj.text or "").strip()
            if not text:
                continue
            style_name = obj.style.name if obj.style is not None else ""
            if style_name.startswith("Heading"):
                buf.append(f"<h3 style='margin:10px 0 6px;color:#1f3d2a;'>{html.escape(text)}</h3>")
            else:
                buf.append(f"<p style='margin:0 0 8px;line-height:1.35;'>{html.escape(text)}</p>")
            if sum(len(x) for x in buf) > 3500:
                flush()
        elif kind == "tbl":
            table_html = _table_to_html(obj)
            if table_html:
                # Keep table in its own chunk so it is not split mid-structure
                flush()
                chunks.append(table_html)
    flush()
    return chunks or ["<p>[No extractable content]</p>"]


def _append_html_pages(pdf, title: str, html_chunks: list[str], fontfile: Path | None) -> None:
    import fitz

    header = _clean_text(title or "Section")
    for i, chunk in enumerate(html_chunks):
        page = pdf.new_page(width=595, height=842)
        _pdf_insert_text(
            page,
            fitz.Rect(48, 36, 547, 64),
            header if i == 0 else f"{header} (cont.)",
            fontsize=13,
            fontfile=fontfile,
            color=(0.12, 0.35, 0.18),
        )
        rect = fitz.Rect(48, 72, 547, 800)
        try:
            page.insert_htmlbox(rect, chunk)
        except Exception:
            # Fallback to plain text extraction from the html chunk
            plain = re.sub(r"<[^>]+>", " ", chunk)
            plain = html.unescape(plain)
            _pdf_insert_text(page, rect, plain, fontsize=10, fontfile=fontfile)


def _append_text_pages(pdf, title: str, body: str, fontfile: Path | None) -> None:
    import fitz

    header = _clean_text(title or "Section")
    text = _clean_text(body or "").strip() or "[No extractable text]"
    max_chars = 3000
    chunks = [text[i : i + max_chars] for i in range(0, len(text), max_chars)] or [text]
    for i, chunk in enumerate(chunks):
        page = pdf.new_page(width=595, height=842)
        _pdf_insert_text(
            page,
            fitz.Rect(48, 40, 547, 70),
            header if i == 0 else f"{header} (cont.)",
            fontsize=14,
            fontfile=fontfile,
            color=(0.12, 0.35, 0.18),
        )
        _pdf_insert_text(
            page,
            fitz.Rect(48, 80, 547, 800),
            chunk,
            fontsize=10,
            fontfile=fontfile,
            color=(0.1, 0.1, 0.1),
        )


def build_compiled_pdf(sources: list[dict]) -> bytes:
    import fitz

    fontfile = _resolve_unicode_font()
    pdf = fitz.open()
    cover = pdf.new_page(width=595, height=842)
    _pdf_insert_text(
        cover,
        fitz.Rect(60, 200, 535, 280),
        "Kapeng Barako - Compiled GI Documentation",
        fontsize=20,
        fontfile=fontfile,
        color=(0.1, 0.3, 0.15),
        align=1,
    )
    _pdf_insert_text(
        cover,
        fitz.Rect(60, 300, 535, 400),
        (
            f"Beanthentic IPOPHL compile\n"
            f"{datetime.now().strftime('%Y-%m-%d %H:%M')}\n"
            f"{len(sources)} source file(s) from Phases 1-3"
        ),
        fontsize=11,
        fontfile=fontfile,
        color=(0.35, 0.35, 0.35),
        align=1,
    )

    for source in sources:
        path: Path = source["path"]
        label = _clean_text(str(source.get("label") or "Document"))
        fname = _clean_text(str(source.get("original_filename") or path.name))
        suffix = path.suffix.lower()

        div = pdf.new_page(width=595, height=842)
        _pdf_insert_text(
            div,
            fitz.Rect(60, 280, 535, 360),
            label,
            fontsize=18,
            fontfile=fontfile,
            color=(0.12, 0.35, 0.18),
            align=1,
        )
        _pdf_insert_text(
            div,
            fitz.Rect(60, 370, 535, 430),
            fname,
            fontsize=10,
            fontfile=fontfile,
            color=(0.4, 0.4, 0.4),
            align=1,
        )

        if suffix == ".pdf":
            try:
                with fitz.open(str(path)) as src:
                    pdf.insert_pdf(src)
                continue
            except Exception:
                body = _extract_text_from_pdf(path) if path.exists() else ""
                _append_text_pages(pdf, label, body or f"[Could not read PDF: {fname}]", fontfile)
                continue

        if suffix in _IMAGE_SUFFIXES:
            page = pdf.new_page(width=595, height=842)
            try:
                page.insert_image(fitz.Rect(40, 60, 555, 780), filename=str(path))
            except Exception:
                _pdf_insert_text(
                    page,
                    fitz.Rect(60, 300, 535, 400),
                    f"[Could not embed image: {fname}]",
                    fontsize=12,
                    fontfile=fontfile,
                    align=1,
                )
            continue

        if suffix == ".docx":
            try:
                chunks = _docx_to_html_chunks(path)
                _append_html_pages(pdf, label, chunks, fontfile)
                continue
            except Exception:
                pass

        body = _source_body_text(path)
        _append_text_pages(pdf, label, body, fontfile)

    out = io.BytesIO()
    pdf.save(out)
    pdf.close()
    return out.getvalue()


def compile_ipophl_package(
    *,
    fmt: str = "pdf",
    file_uuids: list[str] | None = None,
) -> tuple[bytes, str, str, list[dict]]:
    """
    Build compiled package bytes.
    Returns (payload, download_name, mimetype, source_summaries).
    """
    fmt_norm = str(fmt or "pdf").strip().lower()
    if fmt_norm not in {"pdf", "docx"}:
        raise ValueError("format must be 'pdf' or 'docx'")

    sources = collect_compile_sources(file_uuids=file_uuids)
    if not sources:
        raise ValueError(
            "No Phase 1-3 documents found to compile. Upload files in Justification, "
            "Technical, and Control & Traceability first."
        )

    stamp = datetime.now().strftime("%Y%m%d")
    if fmt_norm == "docx":
        payload = build_compiled_docx(sources)
        name = f"Kapeng_Barako_GI_Compiled_{stamp}.docx"
        mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    else:
        payload = build_compiled_pdf(sources)
        name = f"Kapeng_Barako_GI_Compiled_{stamp}.pdf"
        mime = "application/pdf"

    summaries = [
        {
            "file_uuid": s["file_uuid"],
            "task_id": s["task_id"],
            "label": s["label"],
            "original_filename": s["original_filename"],
        }
        for s in sources
    ]
    return payload, name, mime, summaries
